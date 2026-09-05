# -*- coding: utf-8 -*-
"""Convert the Blindband submission markdown into a Google-Docs-ready .docx.

Handles only the subset SUBMISSION.md actually uses: ATX headings, paragraphs,
pipe tables (including cells that hold nothing but an image), fenced code
blocks, images, ordered and unordered lists, horizontal rules, and the inline
run of **bold**, `code` and [links](url).
"""
import os
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SRC = sys.argv[1]
OUT = sys.argv[2]
ROOT = os.path.dirname(os.path.abspath(SRC))

PAGE_W = 6.5  # inches of text column on US Letter with 1" margins
INK = RGBColor(0x1A, 0x1C, 0x20)
QUIET = RGBColor(0x4C, 0x53, 0x5B)
CODE_BG = "F4F2EE"
RULE = RGBColor(0xB9, 0xB2, 0xA4)

IMG_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)$")
INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|https?://[^\s)\]]+)")
LINK_RE = re.compile(r"^\[([^\]]+)\]\(([^)]+)\)$")
BARE_URL_RE = re.compile(r"^https?://[^\s)\]]+$")

# A repository-relative link is dead the moment the document leaves the repo,
# so it is rewritten to the file on GitHub rather than dropped.
REPO_BLOB = "https://github.com/bryankwandou/blindband/blob/main/"


def resolve(url):
    if url.startswith(("http://", "https://", "mailto:")):
        return url
    path = os.path.normpath(os.path.join(ROOT, url)).replace("\\", "/")
    marker = "blindband/"
    i = path.lower().rfind(marker)
    return REPO_BLOB + path[i + len(marker):] if i >= 0 else url


def shade(el, fill):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), fill)
    el.append(sh)


def add_hyperlink(par, text, url):
    part = par.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "8F6210")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(u)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    link.append(run)
    par._p.append(link)


def write_inline(par, text, size=10.5, color=INK, bold=False):
    """Render one line of markdown inline syntax into an existing paragraph."""
    for piece in INLINE_RE.split(text):
        if not piece:
            continue
        m = LINK_RE.match(piece)
        if m:
            add_hyperlink(par, m.group(1).replace("`", ""), resolve(m.group(2)))
            continue
        if BARE_URL_RE.match(piece):
            add_hyperlink(par, piece, piece)
            continue
        if piece.startswith("**") and piece.endswith("**"):
            r = par.add_run(piece[2:-2])
            r.bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            r = par.add_run(piece[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(size - 1)
            r.font.color.rgb = QUIET
            continue
        else:
            r = par.add_run(piece)
            r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = color


def add_image(container, path, alt, width):
    full = os.path.normpath(os.path.join(ROOT, path))
    if not os.path.exists(full):
        print("  ! missing image:", path)
        return
    par = container.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(full, width=Inches(width))
    if alt:
        cap = container.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(alt)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = QUIET


def add_code(doc, lines):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(6)
    par.paragraph_format.space_after = Pt(10)
    par.paragraph_format.left_indent = Inches(0.12)
    shade(par._p.get_or_add_pPr(), CODE_BG)
    for i, line in enumerate(lines):
        if i:
            par.add_run().add_break()
        r = par.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
        r.font.color.rgb = INK


def add_rule(doc):
    par = doc.add_paragraph()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "B9B2A4")
    pbdr.append(bottom)
    par._p.get_or_add_pPr().append(pbdr)


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def add_table(doc, rows):
    head, body = rows[0], rows[1:]
    cols = len(head)
    table = doc.add_table(rows=0, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    width = PAGE_W / cols

    header_is_blank = all(c == "" for c in head)
    source = body if header_is_blank else rows

    for ri, cells in enumerate(source):
        row = table.add_row()
        for ci in range(cols):
            cell = row.cells[ci]
            cell.width = Inches(width)
            text = cells[ci] if ci < len(cells) else ""
            par = cell.paragraphs[0]
            m = IMG_RE.match(text)
            if m:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                full = os.path.normpath(os.path.join(ROOT, m.group(2)))
                if os.path.exists(full):
                    par.add_run().add_picture(full, width=Inches(width - 0.15))
                continue
            is_head = ri == 0 and not header_is_blank
            write_inline(par, text, size=9.5, bold=is_head)
            if is_head:
                shade(cell._tc.get_or_add_tcPr(), CODE_BG)
    return table


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, m, Inches(1))

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    lines = open(SRC, encoding="utf-8").read().split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith("```"):
            block, i = [], i + 1
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i].rstrip())
                i += 1
            add_code(doc, block)
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and set(lines[i + 1].replace("|", "").replace(" ", "")) <= set("-:"):
            rows = [split_row(line)]
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            add_table(doc, rows)
            doc.add_paragraph()
            continue

        if not line.strip():
            i += 1
            continue

        if line.strip() in ("---", "***", "___"):
            add_rule(doc)
            i += 1
            continue

        m = IMG_RE.match(line.strip())
        if m:
            add_image(doc, m.group(2), m.group(1), PAGE_W)
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
            i += 1
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
            i += 1
            continue

        bullet = re.match(r"^([-*])\s+(.*)$", line)
        number = re.match(r"^(\d+)\.\s+(.*)$", line)
        if bullet or number:
            text = (bullet or number).group(2)
            # A list item may wrap onto continuation lines indented under it.
            while i + 1 < len(lines) and lines[i + 1].startswith("   ") and lines[i + 1].strip():
                i += 1
                text += " " + lines[i].strip()
            style = "List Bullet" if bullet else "List Number"
            par = doc.add_paragraph(style=style)
            write_inline(par, text)
            i += 1
            continue

        # Plain paragraph: join the wrapped lines that belong to it.
        buf = [line.strip()]
        while i + 1 < len(lines):
            nxt = lines[i + 1].rstrip()
            if (
                not nxt.strip()
                or nxt.startswith(("#", "|", "```", "- ", "* ", "!["))
                or re.match(r"^\d+\.\s", nxt)
                or nxt.strip() in ("---", "***", "___")
            ):
                break
            i += 1
            buf.append(nxt.strip())
        par = doc.add_paragraph()
        write_inline(par, " ".join(buf))
        i += 1

    doc.save(OUT)
    print("wrote", OUT)


main()
